"""
Report Generator - Generate personalized resume reports using Ollama
Standalone module - does not affect core functionality
"""
import requests
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List
import json

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL = "llama3.2:3b"

def generate_personalized_recommendations(resume_data: Dict, analysis_data: Dict) -> List[str]:
    """Generate personalized recommendations using Ollama"""
    
    score = analysis_data.get('final_score', 0)
    breakdown = analysis_data.get('breakdown', {})
    
    # Find top 3 weakest areas
    weak_areas = sorted(breakdown.items(), key=lambda x: x[1])[:3] if breakdown else []
    weak_areas_text = ", ".join([f"{k.replace('_', ' ')} ({v:.0f}/100)" for k, v in weak_areas]) if weak_areas else "N/A"
    
    prompt = f"""You are an ATS resume expert. Generate 5 specific, actionable recommendations for this candidate.

Candidate: {resume_data.get('name', 'N/A')}
Current Score: {score}/100
Top Skills: {', '.join(resume_data.get('skills', [])[:5])}

Weakest Areas:
{weak_areas_text}

Generate EXACTLY 5 recommendations in this format:
1. [Specific action for weakest area]
2. [Specific action for second weak area]
3. [Specific action for third weak area]
4. [General improvement suggestion]
5. [Career advancement tip]

Be specific, actionable, and professional. Each recommendation should be 1-2 sentences."""

    try:
        response = requests.post(
            OLLAMA_URL,
            json={
                "model": MODEL,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.7, "top_p": 0.9}
            },
            timeout=30
        )
        
        if response.status_code == 200:
            content = response.json().get('response', '')
            recs = []
            for line in content.split('\n'):
                line = line.strip()
                if line and (line[0].isdigit() or line.startswith('-') or line.startswith('•')):
                    rec = line.lstrip('0123456789.-•) ').strip()
                    if rec and len(rec) > 10:
                        recs.append(rec)
            
            if len(recs) >= 3:
                return recs[:5]
    except Exception as e:
        print(f"Ollama recommendations error: {e}")
    
    # Fallback to template-based
    feedback = analysis_data.get('feedback', [])
    if feedback:
        return feedback[:5]
    return [
        "Add quantifiable achievements with metrics (%, $, numbers)",
        "Tailor resume keywords to match job requirements",
        "Demonstrate skills through specific project examples",
        "Use stronger action verbs to show impact",
        "Enhance document formatting and structure"
    ]

def generate_report_content(resume_data: Dict, analysis_data: Dict, resume_text: str = "", jd_text: str = "") -> str:
    """Generate detailed report using Ollama with full context"""
    
    score = analysis_data.get('final_score', 0)
    grade = analysis_data.get('grade', 'N/A')
    breakdown = analysis_data.get('breakdown', {})
    
    # Get matched/missing skills
    matched_skills = analysis_data.get('skill_match_details', {}).get('matched', [])
    missing_skills = analysis_data.get('skill_match_details', {}).get('missing', [])
    
    # Find weakest areas
    weak_areas = sorted(breakdown.items(), key=lambda x: x[1])[:3]
    strong_areas = sorted(breakdown.items(), key=lambda x: x[1], reverse=True)[:3]
    
    # Truncate texts for prompt
    resume_snippet = resume_text[:800] if resume_text else resume_data.get('name', 'N/A')
    jd_snippet = jd_text[:600] if jd_text else "General position"
    
    prompt = f"""You are an ATS resume expert. Analyze this candidate against the job requirements.

JOB DESCRIPTION:
{jd_snippet}

RESUME EXCERPT:
{resume_snippet}

ANALYSIS SCORES:
- Overall Score: {score}/100 (Grade: {grade})
- Strongest Areas: {', '.join([f"{k.replace('_', ' ')}: {v:.0f}" for k,v in strong_areas])}
- Weakest Areas: {', '.join([f"{k.replace('_', ' ')}: {v:.0f}" for k,v in weak_areas])}
- Matched Skills: {', '.join(matched_skills[:8]) if matched_skills else 'None identified'}
- Missing Skills: {', '.join(missing_skills[:8]) if missing_skills else 'None identified'}

Generate a report in this EXACT format:

EXECUTIVE SUMMARY
[2-3 sentences analyzing how well this candidate fits the role based on the JD requirements and their resume. Mention specific skills and scores.]

STRENGTHS
• [Specific strength from resume that matches JD requirement]
• [Another strength with evidence from their experience]
• [Third strength highlighting their best scoring area]

AREAS FOR IMPROVEMENT
• [Specific weakness from lowest scoring area with actionable fix]
• [Missing skill from JD with suggestion to acquire it]
• [Another improvement area with concrete action]

NEXT STEPS
1. [Actionable step addressing weakest area]
2. [Actionable step for missing critical skill]
3. [Actionable step for overall improvement]

Be specific, reference actual skills/requirements, and provide actionable advice."""

    try:
        response = requests.post(
            OLLAMA_URL,
            json={
                "model": MODEL, 
                "prompt": prompt, 
                "stream": False,
                "options": {
                    "temperature": 0.3,
                    "top_p": 0.9,
                    "repeat_penalty": 1.1
                }
            },
            timeout=45
        )
        
        if response.status_code == 200:
            content = response.json().get('response', '')
            if len(content) > 100 and 'EXECUTIVE SUMMARY' in content:
                return content
            else:
                return generate_fallback_report(resume_data, analysis_data, resume_text, jd_text)
        else:
            return generate_fallback_report(resume_data, analysis_data, resume_text, jd_text)
    except Exception as e:
        print(f"Ollama error: {e}")
        return generate_fallback_report(resume_data, analysis_data, resume_text, jd_text)

def generate_fallback_report(resume_data: Dict, analysis_data: Dict, resume_text: str = "", jd_text: str = "") -> str:
    """Fallback report if Ollama unavailable"""
    score = analysis_data.get('final_score', 0)
    grade = analysis_data.get('grade', 'N/A')
    breakdown = analysis_data.get('breakdown', {})
    
    matched_skills = analysis_data.get('skill_match_details', {}).get('matched', [])
    missing_skills = analysis_data.get('skill_match_details', {}).get('missing', [])
    
    weak_areas = sorted(breakdown.items(), key=lambda x: x[1])[:3]
    strong_areas = sorted(breakdown.items(), key=lambda x: x[1], reverse=True)[:3]
    
    return f"""EXECUTIVE SUMMARY
Your resume scored {score}/100 (Grade {grade}). Analysis shows {len(matched_skills)} matched skills and {len(missing_skills)} missing skills from the job requirements. Your strongest areas are {', '.join([k.replace('_', ' ') for k,v in strong_areas[:2]])}.

STRENGTHS
• {strong_areas[0][0].replace('_', ' ').title()}: Scored {strong_areas[0][1]:.0f}/100
• Matched {len(matched_skills)} required skills: {', '.join(matched_skills[:5]) if matched_skills else 'Basic alignment'}
• {strong_areas[1][0].replace('_', ' ').title() if len(strong_areas) > 1 else 'Professional presentation'}

AREAS FOR IMPROVEMENT
• {weak_areas[0][0].replace('_', ' ').title()}: Scored only {weak_areas[0][1]:.0f}/100 - needs significant improvement
• Missing critical skills: {', '.join(missing_skills[:5]) if missing_skills else 'Add more job-specific keywords'}
• {weak_areas[1][0].replace('_', ' ').title() if len(weak_areas) > 1 else 'Quantify achievements with metrics'}

NEXT STEPS
1. Focus on improving {weak_areas[0][0].replace('_', ' ')} (lowest scoring area)
2. Acquire or highlight these missing skills: {', '.join(missing_skills[:3]) if missing_skills else 'job-specific keywords'}
3. Add measurable achievements with numbers, percentages, or dollar amounts"""

def create_personalized_report(resume_data: Dict, analysis_data: Dict, resume_text: str, jd_text: str, output_path: str):
    """Create personalized report using resume text, JD text, and scores"""
    from .personalized_report_generator import create_personalized_report as gen_report
    gen_report(resume_data, analysis_data, resume_text, jd_text, output_path)

def create_docx_report_with_context(resume_data: Dict, analysis_data: Dict, resume_text: str, jd_text: str, output_path: str):
    """Create DOCX report with full context for AI generation"""
    # Store context in analysis_data for use in create_docx_report
    analysis_data['_resume_text'] = resume_text
    analysis_data['_jd_text'] = jd_text
    create_docx_report(resume_data, analysis_data, output_path)

def create_docx_report(resume_data: Dict, analysis_data: Dict, output_path: str):
    """Create formatted DOCX report"""
    
    doc = Document()
    
    # Header
    header = doc.add_heading('ATS Resume Analysis Report', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Candidate Info (FIRST)
    doc.add_heading('Candidate Information', 1)
    p = doc.add_paragraph()
    p.add_run(f"Name: ").bold = True
    p.add_run(f"{resume_data.get('name', 'N/A')}\n")
    p.add_run(f"Email: ").bold = True
    p.add_run(f"{resume_data.get('email', 'N/A')}\n")
    p.add_run(f"Analysis Date: ").bold = True
    from datetime import datetime
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
    
    doc.add_paragraph()
    
    # AI-Generated Personalized Recommendations (SECOND - Most Actionable)
    doc.add_heading('Personalized Recommendations', 1)
    personalized_recs = generate_personalized_recommendations(resume_data, analysis_data)
    for rec in personalized_recs:
        doc.add_paragraph(rec, style='List Bullet')
    
    doc.add_paragraph()
    
    # AI-Generated Report (THIRD)
    doc.add_heading('AI-Generated Insights', 1)
    resume_text = analysis_data.get('_resume_text', '')
    jd_text = analysis_data.get('_jd_text', '')
    report_content = generate_report_content(resume_data, analysis_data, resume_text, jd_text)
    for line in report_content.split('\n'):
        if line.strip():
            if line.isupper() and len(line) < 50:
                doc.add_heading(line.strip(), 2)
            else:
                doc.add_paragraph(line.strip())
    
    doc.add_paragraph()
    
    # Top Deductions (FOURTH)
    doc.add_heading('Highest Point Deductions', 1)
    breakdown = analysis_data.get('breakdown', {})
    
    categories = {
        'file_layout': {'weight': 10, 'desc': 'Document structure, sections, formatting'},
        'font_consistency': {'weight': 5, 'desc': 'Font uniformity and readability'},
        'readability': {'weight': 5, 'desc': 'Language clarity and grammar'},
        'professional_language': {'weight': 10, 'desc': 'Action verbs and professional tone'},
        'date_consistency': {'weight': 2.5, 'desc': 'Date format and chronology'},
        'employment_gaps': {'weight': 5, 'desc': 'Career continuity'},
        'career_progression': {'weight': 2.5, 'desc': 'Role advancement and growth'},
        'keyword_alignment': {'weight': 15, 'desc': 'Job description keyword match'},
        'skill_context': {'weight': 5, 'desc': 'Skills demonstrated in experience'},
        'semantic_fit': {'weight': 25, 'desc': 'Overall job-role alignment (AI-based)'},
        'quantified_impact': {'weight': 15, 'desc': 'Measurable achievements (%, $, numbers)'},
        'online_presence': {'weight': 5, 'desc': 'LinkedIn and portfolio links'}
    }
    
    deductions = []
    for key, value in breakdown.items():
        if key in categories:
            points_lost = categories[key]['weight'] * (1 - value/100)
            if points_lost > 0.5:
                deductions.append((key, points_lost, value))
    
    deductions.sort(key=lambda x: x[1], reverse=True)
    for key, lost, score in deductions[:5]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
        p.add_run(f"-{lost:.2f} points (scored {score:.1f}/100)")
    
    doc.add_paragraph()
    
    # Score Impact Summary (FIFTH)
    doc.add_heading('Score Impact Summary', 1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    total_points = sum((breakdown.get(k, 0) / 100) * categories[k]['weight'] for k in categories if k in breakdown)
    total_possible = sum(categories[k]['weight'] for k in categories)
    total_deducted = total_possible - total_points
    
    summary_table.cell(0, 0).text = 'Total Points Earned'
    summary_table.cell(0, 1).text = f"{total_points:.2f}"
    summary_table.cell(1, 0).text = 'Total Possible'
    summary_table.cell(1, 1).text = f"{total_possible:.2f}"
    summary_table.cell(2, 0).text = 'Points Deducted'
    summary_table.cell(2, 1).text = f"{total_deducted:.2f}"
    summary_table.cell(3, 0).text = 'Final Percentage'
    summary_table.cell(3, 1).text = f"{analysis_data.get('final_score', 0):.1f}%"
    
    doc.add_paragraph()
    
    # Detailed Score Analysis (SIXTH)
    doc.add_heading('Detailed Score Analysis', 1)
    for key, value in breakdown.items():
        if key in categories:
            cat_info = categories[key]
            
            p = doc.add_paragraph()
            p.add_run(f"{key.replace('_', ' ').title()}").bold = True
            p.add_run(f" ({cat_info['weight']}% weight)\n")
            
            score_text = f"Score: {value:.1f}/100 "
            if value >= 80:
                score_text += "[Excellent]"
            elif value >= 60:
                score_text += "[Good]"
            elif value >= 40:
                score_text += "[Needs Improvement]"
            else:
                score_text += "[Critical]"
            
            p.add_run(score_text + "\n")
            p.add_run(f"Impact: {cat_info['desc']}\n").italic = True
            
            points_earned = (value / 100) * cat_info['weight']
            points_possible = cat_info['weight']
            points_lost = points_possible - points_earned
            
            p.add_run(f"Points: {points_earned:.2f}/{points_possible} ")
            if points_lost > 0.5:
                p.add_run(f"(-{points_lost:.2f} deducted)").font.color.rgb = RGBColor(220, 38, 38)
            
            doc.add_paragraph()
    
    # Score Summary (SEVENTH - LAST)
    doc.add_heading('Score Summary', 1)
    score_table = doc.add_table(rows=2, cols=2)
    score_table.style = 'Light Grid Accent 1'
    
    score_table.cell(0, 0).text = 'Overall Score'
    score_table.cell(0, 1).text = f"{analysis_data.get('final_score', 0)}/100"
    score_table.cell(1, 0).text = 'Grade'
    score_table.cell(1, 1).text = analysis_data.get('grade', 'N/A')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Generated by ATS Intelligence Engine')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save(output_path)
