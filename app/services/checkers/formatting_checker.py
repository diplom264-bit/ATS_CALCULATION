"""
Phase 1: Formatting & ATS-Compliance Module
Checks file layout, fonts, colors for ATS compatibility
"""
import fitz  # PyMuPDF
from docx import Document
from typing import Dict, List, Tuple

class FormattingChecker:
    """Rule-based file analysis for ATS compatibility"""
    
    STANDARD_FONTS = {'Arial', 'Calibri', 'Times New Roman', 'Helvetica', 'Georgia'}
    
    def check_file_layout(self, file_path: str) -> Tuple[float, List[str]]:
        """Check 1: File Type & Layout (20 points)"""
        score = 20.0
        feedback = []
        
        if file_path.endswith('.pdf'):
            score_delta, msgs = self._check_pdf_layout(file_path)
            score += score_delta
            feedback.extend(msgs)
        elif file_path.endswith('.docx'):
            score_delta, msgs = self._check_docx_layout(file_path)
            score += score_delta
            feedback.extend(msgs)
        elif file_path.endswith('.txt'):
            score -= 10
            feedback.append("TXT format lacks formatting - use PDF or DOCX")
        
        return max(0, score), feedback
    
    def _check_pdf_layout(self, file_path: str) -> Tuple[float, List[str]]:
        """Check PDF for tables and images"""
        score = 0.0
        feedback = []
        
        try:
            doc = fitz.open(file_path)
            for page in doc:
                # Check tables
                tables = page.find_tables()
                if tables:
                    score -= 5
                    feedback.append("Tables detected - may break ATS parsers")
                
                # Check images
                images = page.get_images()
                if images:
                    score -= 5
                    feedback.append("Images/graphics detected - avoid visual elements")
            doc.close()
        except:
            pass
        
        return score, feedback
    
    def _check_docx_layout(self, file_path: str) -> Tuple[float, List[str]]:
        """Check DOCX for tables and images"""
        score = 0.0
        feedback = []
        
        try:
            doc = Document(file_path)
            if len(doc.tables) > 0:
                score -= 5
                feedback.append(f"{len(doc.tables)} tables found - use simple formatting")
            
            image_count = sum(len(run.element.xpath('.//pic:pic')) for para in doc.paragraphs for run in para.runs)
            if image_count > 0:
                score -= 5
                feedback.append("Images detected - remove visual elements")
        except:
            pass
        
        return score, feedback
    
    def check_font_consistency(self, file_path: str) -> Tuple[float, List[str]]:
        """Check 2: Font & Color Consistency (10 points)"""
        score = 10.0
        feedback = []
        
        if file_path.endswith('.pdf'):
            score_delta, msgs = self._check_pdf_fonts(file_path)
            score += score_delta
            feedback.extend(msgs)
        elif file_path.endswith('.docx'):
            score_delta, msgs = self._check_docx_fonts(file_path)
            score += score_delta
            feedback.extend(msgs)
        
        return max(0, score), feedback
    
    def _check_pdf_fonts(self, file_path: str) -> Tuple[float, List[str]]:
        """Check PDF fonts, sizes, colors"""
        score = 0.0
        feedback = []
        fonts = set()
        sizes = []
        
        try:
            doc = fitz.open(file_path)
            for page in doc:
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                fonts.add(span["font"])
                                sizes.append(span["size"])
            doc.close()
            
            # Check font variety
            if len(fonts) > 3:
                score -= 3
                feedback.append(f"Too many fonts ({len(fonts)}) - use 1-2 standard fonts")
            
            # Check font types
            non_standard = [f for f in fonts if not any(sf in f for sf in self.STANDARD_FONTS)]
            if non_standard:
                score -= 2
                feedback.append("Use standard fonts (Arial, Calibri, Times New Roman)")
            
            # Check size consistency
            body_sizes = [s for s in sizes if 9 <= s <= 13]
            if body_sizes and (min(body_sizes) < 10 or max(body_sizes) > 12):
                score -= 2
                feedback.append("Use 10-12pt font for body text")
        except:
            pass
        
        return score, feedback
    
    def _check_docx_fonts(self, file_path: str) -> Tuple[float, List[str]]:
        """Check DOCX fonts and sizes"""
        score = 0.0
        feedback = []
        fonts = set()
        sizes = []
        
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.name:
                        fonts.add(run.font.name)
                    if run.font.size:
                        sizes.append(run.font.size.pt)
            
            if len(fonts) > 3:
                score -= 3
                feedback.append(f"Too many fonts ({len(fonts)}) - simplify formatting")
            
            non_standard = fonts - self.STANDARD_FONTS
            if non_standard:
                score -= 2
                feedback.append("Use standard fonts")
            
            body_sizes = [s for s in sizes if 9 <= s <= 13]
            if body_sizes and (min(body_sizes) < 10 or max(body_sizes) > 12):
                score -= 2
                feedback.append("Use 10-12pt font size")
        except:
            pass
        
        return score, feedback
