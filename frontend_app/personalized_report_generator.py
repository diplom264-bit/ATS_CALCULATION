"""
Personalized Report Generator
Generates truly adaptive recommendations based on individual scores and context
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List
import re

def create_personalized_report(resume_data: Dict, analysis_data: Dict, resume_text: str, jd_text: str, output_path: str):
    """Generate personalized DOCX report with adaptive recommendations"""
    doc = Document()
    
    # Title
    title = doc.add_heading('ATS Resume Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Candidate Info
    doc.add_heading('Candidate Information', 1)
    doc.add_paragraph(f"Name: {resume_data.get('name', 'N/A')}")
    doc.add_paragraph(f"Email: {resume_data.get('email', 'N/A')}")
    doc.add_paragraph(f"Score: {analysis_data.get('final_score', 0):.1f}/100 (Grade {analysis_data.get('grade', 'F')})")
    
    # Score Breakdown
    doc.add_heading('Score Breakdown', 1)
    breakdown = analysis_data.get('breakdown', {})
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Score'
    
    for key, value in breakdown.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key.replace('_', ' ').title()
        row_cells[1].text = f"{value:.1f}/100"
    
    # Personalized Recommendations
    doc.add_heading('Personalized Recommendations', 1)
    recommendations = generate_personalized_recommendations(
        breakdown, resume_data, resume_text, jd_text, analysis_data
    )
    
    for i, rec in enumerate(recommendations, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(f"{rec['title']}\n").bold = True
        p.add_run(rec['detail'])
        if rec.get('action'):
            p.add_run(f"\nAction: ").bold = True
            p.add_run(rec['action'])
    
    doc.save(output_path)

def generate_personalized_recommendations(
    breakdown: Dict, 
    resume_data: Dict, 
    resume_text: str, 
    jd_text: str,
    analysis_data: Dict
) -> List[Dict]:
    """Generate truly personalized recommendations using Ollama Llama 2 3B"""
    
    # Try Ollama first for personalized recommendations
    ollama_recs = _generate_ollama_recommendations(breakdown, resume_data, resume_text, jd_text, analysis_data)
    if ollama_recs:
        return ollama_recs
    
    # Fallback to rule-based
    recommendations = []
    sorted_scores = sorted(breakdown.items(), key=lambda x: x[1])[:3]
    
    for category, score in sorted_scores:
        if score < 70:
            rec = _generate_category_recommendation(
                category, score, resume_data, resume_text, jd_text, breakdown, analysis_data
            )
            if rec:
                recommendations.append(rec)
    
    if analysis_data.get('final_score', 0) < 50:
        recommendations.insert(0, {
            'title': '🚨 Critical: Role Mismatch Detected',
            'detail': 'Your resume shows significant misalignment with the job requirements. Consider whether this role truly matches your experience.',
            'action': 'Review the job description carefully and either tailor your resume to highlight relevant experience, or consider applying to roles that better match your background.'
        })
    
    return recommendations[:5]

def _generate_ollama_recommendations(
    breakdown: Dict,
    resume_data: Dict,
    resume_text: str,
    jd_text: str,
    analysis_data: Dict
) -> List[Dict]:
    """Generate recommendations using Ollama Llama 2 3B"""
    try:
        import requests
        
        # Get weakest areas
        sorted_scores = sorted(breakdown.items(), key=lambda x: x[1])[:3]
        weak_areas = ', '.join([f"{k.replace('_', ' ')} ({v:.0f}/100)" for k, v in sorted_scores])
        
        # Get missing skills
        skill_details = analysis_data.get('skill_match_details', {})
        missing_skills = ', '.join(skill_details.get('missing', [])[:5])
        
        prompt = f"""You are an ATS resume expert. Generate 5 specific, actionable recommendations for this candidate.

Candidate: {resume_data.get('name', 'N/A')}
Current Score: {analysis_data.get('final_score', 0):.0f}/100
Weakest Areas: {weak_areas}
Missing Skills: {missing_skills}

Generate EXACTLY 5 recommendations in this format:
1. [Title] - [Specific action with details]
2. [Title] - [Specific action with details]
3. [Title] - [Specific action with details]
4. [Title] - [Specific action with details]
5. [Title] - [Specific action with details]

Be specific, actionable, and professional. Each recommendation should be 2-3 sentences."""

        response = requests.post(
            'http://localhost:11434/api/generate',
            json={
                'model': 'llama2:3b',
                'prompt': prompt,
                'stream': False,
                'options': {'temperature': 0.7, 'top_p': 0.9}
            },
            timeout=30
        )
        
        if response.status_code == 200:
            content = response.json().get('response', '')
            recs = []
            
            for line in content.split('\n'):
                line = line.strip()
                if line and (line[0].isdigit() or line.startswith('-')):
                    # Parse "Title - Detail" format
                    parts = line.lstrip('0123456789.-) ').split(' - ', 1)
                    if len(parts) == 2:
                        recs.append({
                            'title': parts[0].strip(),
                            'detail': parts[1].strip(),
                            'action': ''
                        })
                    elif len(parts[0]) > 10:
                        recs.append({
                            'title': 'Recommendation',
                            'detail': parts[0].strip(),
                            'action': ''
                        })
            
            if len(recs) >= 3:
                return recs[:5]
    except Exception as e:
        import logging
        logging.warning(f"Ollama recommendations failed: {e}")
    
    return None

def _generate_category_recommendation(
    category: str, 
    score: float, 
    resume_data: Dict, 
    resume_text: str, 
    jd_text: str,
    breakdown: Dict,
    analysis_data: Dict
) -> Dict:
    """Generate specific recommendation for a category"""
    
    if category == 'semantic_fit':
        # Extract missing keywords from JD
        jd_keywords = _extract_keywords(jd_text)
        resume_keywords = _extract_keywords(resume_text)
        missing = [k for k in jd_keywords[:10] if k not in resume_keywords]
        
        return {
            'title': f'Semantic Job-Fit ({score:.0f}/100)',
            'detail': f'Your resume shows weak alignment with the job requirements. Missing key terms: {", ".join(missing[:5])}.',
            'action': f'Add these specific skills/terms to your resume: {", ".join(missing[:3])}. Ensure they appear in context within your experience section.'
        }
    
    elif category == 'keyword_alignment':
        skill_details = analysis_data.get('skill_match_details', {})
        missing_skills = skill_details.get('missing', [])[:5]
        
        return {
            'title': f'Keyword Alignment ({score:.0f}/100)',
            'detail': f'Your resume is missing {len(missing_skills)} critical keywords from the job description.',
            'action': f'Add these keywords naturally: {", ".join(missing_skills)}. Place them in your skills section and demonstrate them in your experience bullets.'
        }
    
    elif category == 'quantified_impact':
        has_numbers = bool(re.search(r'\d+%|\$\d+|\d+x', resume_text))
        
        return {
            'title': f'Quantified Impact ({score:.0f}/100)',
            'detail': 'Your resume lacks measurable achievements. Recruiters want to see concrete results.' if not has_numbers else 'You have some metrics, but need more quantified achievements.',
            'action': 'Add specific numbers to each bullet: "Increased X by Y%", "Reduced costs by $Z", "Managed team of N people", "Delivered X projects in Y months".'
        }
    
    elif category == 'skill_context':
        skills = resume_data.get('skills', [])
        
        return {
            'title': f'Skill Context ({score:.0f}/100)',
            'detail': f'You list {len(skills)} skills, but many aren\'t demonstrated in your experience section.',
            'action': f'For each skill ({", ".join(skills[:3])}), add a specific example in your work experience showing how you used it to achieve results.'
        }
    
    elif category == 'professional_language':
        return {
            'title': f'Professional Language ({score:.0f}/100)',
            'detail': 'Your resume needs stronger action verbs and professional terminology.',
            'action': 'Replace weak verbs (did, made, worked on) with strong ones (architected, optimized, spearheaded, delivered). Start each bullet with a powerful action verb.'
        }
    
    elif category == 'file_layout':
        return {
            'title': f'File Layout ({score:.0f}/100)',
            'detail': 'Your resume formatting may not be ATS-friendly.',
            'action': 'Use a simple, single-column layout. Avoid tables, text boxes, headers/footers. Use standard section headings: Experience, Education, Skills. Save as .docx or PDF.'
        }
    
    elif category == 'career_progression':
        companies = resume_data.get('companies', [])
        
        return {
            'title': f'Career Progression ({score:.0f}/100)',
            'detail': f'Your career trajectory isn\'t clearly showing growth across {len(companies)} positions.',
            'action': 'Highlight promotions, increased responsibilities, and skill development. Show progression: Junior → Mid → Senior. Emphasize leadership and ownership growth.'
        }
    
    elif category == 'readability':
        word_count = len(resume_text.split())
        
        return {
            'title': f'Readability ({score:.0f}/100)',
            'detail': f'Your resume ({word_count} words) may be too {"long" if word_count > 800 else "short"} or complex.',
            'action': 'Aim for 400-600 words. Use short, punchy bullets (1-2 lines each). Avoid jargon. Write at 8th-grade reading level for ATS parsing.'
        }
    
    elif category == 'online_presence':
        has_linkedin = bool(resume_data.get('linkedin'))
        
        return {
            'title': f'Online Presence ({score:.0f}/100)',
            'detail': 'Missing professional online presence.' if not has_linkedin else 'Limited online visibility.',
            'action': 'Add LinkedIn URL to resume header. Ensure LinkedIn profile is complete (500+ connections, detailed experience, recommendations). Consider GitHub portfolio if technical role.'
        }
    
    return None

def _extract_keywords(text: str) -> List[str]:
    """Extract important keywords from text"""
    # Simple keyword extraction - remove common words
    words = re.findall(r'\b[a-z]{3,}\b', text.lower())
    stop_words = {'the', 'and', 'for', 'with', 'this', 'that', 'from', 'have', 'will', 'your', 'are', 'not', 'but', 'can', 'all', 'one', 'our', 'out', 'use', 'has', 'was', 'were', 'been', 'their', 'they', 'what', 'when', 'where', 'who', 'which', 'how', 'why', 'about', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'between', 'under', 'again', 'further', 'then', 'once'}
    keywords = [w for w in words if w not in stop_words and len(w) > 3]
    
    # Count frequency
    from collections import Counter
    freq = Counter(keywords)
    return [k for k, v in freq.most_common(20)]
